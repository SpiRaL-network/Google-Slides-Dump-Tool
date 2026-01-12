import os
import sys
import subprocess
import re

# ===== 1. Auto-install python-docx =====
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.section import WD_ORIENT
except ImportError:
    print("python-docx not installed. Installing...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches
    from docx.enum.section import WD_ORIENT

# Now that python-docx is available, import low-level stuff
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ===== 2. Auto-install Pillow =====
try:
    from PIL import Image
except ImportError:
    print("Pillow not installed. Installing...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "Pillow"])
    from PIL import Image


# ===== Function: add page number in footer "Page X / Y" =====
def add_page_number(section):
    """
    Add 'Page X / Y' centered inside the footer
    for the given section, using Word fields { PAGE } and { NUMPAGES }.
    """
    footer = section.footer

    # Use first paragraph in footer, or create one
    if footer.paragraphs:
        paragraph = footer.paragraphs[0]
    else:
        paragraph = footer.add_paragraph()

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ----- Field { PAGE } -----
    run_page = paragraph.add_run()

    fld_begin_page = OxmlElement('w:fldChar')
    fld_begin_page.set(qn('w:fldCharType'), 'begin')

    instr_text_page = OxmlElement('w:instrText')
    instr_text_page.set(qn('xml:space'), 'preserve')
    instr_text_page.text = " PAGE "

    fld_separate_page = OxmlElement('w:fldChar')
    fld_separate_page.set(qn('w:fldCharType'), 'separate')

    fld_end_page = OxmlElement('w:fldChar')
    fld_end_page.set(qn('w:fldCharType'), 'end')

    run_page._r.append(fld_begin_page)
    run_page._r.append(instr_text_page)
    run_page._r.append(fld_separate_page)
    run_page._r.append(fld_end_page)

    # Text " / "
    paragraph.add_run(" / ")

    # ----- Field { NUMPAGES } -----
    run_numpages = paragraph.add_run()

    fld_begin_np = OxmlElement('w:fldChar')
    fld_begin_np.set(qn('w:fldCharType'), 'begin')

    instr_text_np = OxmlElement('w:instrText')
    instr_text_np.set(qn('xml:space'), 'preserve')
    instr_text_np.text = " NUMPAGES "

    fld_separate_np = OxmlElement('w:fldChar')
    fld_separate_np.set(qn('w:fldCharType'), 'separate')

    fld_end_np = OxmlElement('w:fldChar')
    fld_end_np.set(qn('w:fldCharType'), 'end')

    run_numpages._r.append(fld_begin_np)
    run_numpages._r.append(instr_text_np)
    run_numpages._r.append(fld_separate_np)
    run_numpages._r.append(fld_end_np)

# ===== 3. Default folder & output =====
folder = "captures"
docx_name = "result.docx"

# ===== 4. Numeric sort for Page_X.xxx =====
print("\n=== Image to DOCX converter ===")
def page_number(filename: str) -> int:
    m = re.match(r"^Page_(\d+)", filename, re.IGNORECASE)
    if m:
        return int(m.group(1))
    return 10**9  # fallback for naming errors


# ===== 5. Collect & sort images =====
print("Loading images...")
extensions = (".png", ".jpg", ".jpeg", ".bmp", ".webp")

try:
    all_files = os.listdir(folder)
except FileNotFoundError:
    print("Folder not found.")
    sys.exit(1)

files = [
    f for f in all_files
    if f.lower().startswith("page_") and f.lower().endswith(extensions)
]

files.sort(key=page_number)

if not files:
    print("No 'Page_X' images found. Files must be named 'Page_X' where 'X' is an INTEGER.")
    sys.exit(1)

page_count = len(files)

# ===== 6. Create DOCX & set LANDSCAPE =====
document = Document()
section = document.sections[0]

# Landscape orientation
section.orientation = WD_ORIENT.LANDSCAPE

# Swap width/height for landscape
page_width = section.page_height
page_height = section.page_width
section.page_width = page_width
section.page_height = page_height

# Set all margins to zero
section.top_margin = Inches(0)
section.bottom_margin = Inches(0)
section.left_margin = Inches(0)
section.right_margin = Inches(0)

# Add "Page X / Y" in footer
add_page_number(section)

# EMUs per inch (python-docx internal unit)
EMU_PER_INCH = 914400

# Compute usable area = entire page
usable_width_emu = int(section.page_width)
usable_height_emu = int(section.page_height)

usable_width_inches = usable_width_emu / EMU_PER_INCH
usable_height_inches = usable_height_emu / EMU_PER_INCH

# ===== 7. Insert each image =====
for file in files:
    path = os.path.join(folder, file)

    img = Image.open(path)
    img_width_px, img_height_px = img.size

    # Assume 96 DPI for screen captures
    img_width_in = img_width_px / 96.0
    img_height_in = img_height_px / 96.0

    img_ratio = img_width_in / img_height_in

    # Fit image to full page width
    final_width_in = usable_width_inches
    final_height_in = final_width_in / img_ratio

    document.add_picture(
        path,
        width=Inches(final_width_in),
        height=Inches(final_height_in)
    )

    document.add_page_break()

print("Creating .DOCX")
document.save(docx_name)

print(f"DOCX created: {docx_name}")
print(f"Total pages: {page_count}")
